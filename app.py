import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# הגדרת כותרת ועיצוב הדף
st.set_page_config(page_title="מחולל פרוטוקולים", page_icon="📝", layout="centered")

# פונקציה שמדברת עם ה-AI
def generate_protocol(api_key, text):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash') # מודל מהיר וחינמי
    
    prompt = f"""
    תפעל כמזכיר ישיבות מקצועי ומנוסה. הנה סיכום גולמי של פגישה:
    "{text}"
    
    אנא צור פרוטוקול פגישה רשמי ומסודר בעברית.
    הקפד על עיצוב נקי ושפה מקצועית.
    המבנה הנדרש:
    1. כותרת: פרוטוקול פגישה
    2. רשימת משתתפים (אם מוזכרים בטקסט)
    3. תקציר הדיון (נקודות עיקריות)
    4. החלטות שהתקבלו
    5. משימות לביצוע (Action Items) - מי עושה מה ומתי.
    
    הפק רק את תוכן הפרוטוקול, ללא טקסט מקדים או מסכם.
    """
    response = model.generate_content(prompt)
    return response.text

# פונקציה שיוצרת קובץ וורד
def create_docx(content):
    doc = Document()
    
    # הגדרת כיוון כללי למסמך (RTL טריק)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    
    title = doc.add_heading('פרוטוקול פגישה', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
    
    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio

# --- הממשק הוויזואלי ---
st.title("כלי ליצירת פרוטוקול פגישה 📝")
st.markdown("---")

# תפריד צד להגדרות
with st.sidebar:
    st.header("הגדרות")
    api_key = st.text_input("הדבק כאן את ה-API Key מגוגל:", type="password")
    st.info("המפתח נשמר רק לזמן הריצה הנוכחי ולא נשמר במערכת.")

# אזור הטקסט הראשי
meeting_notes = st.text_area("הדבק כאן את הסיכום הגולמי של הפגישה:", height=200, placeholder="למשל: דיברנו עם יוסי על הפרויקט החדש, הוחלט שדני יכין מצגת עד יום חמישי...")

if st.button("צור פרוטוקול עכשיו", type="primary"):
    if not api_key:
        st.error("חסר מפתח API. אנא הזן אותו בתפריט הצד.")
    elif not meeting_notes:
        st.warning("לא הזנת תוכן לפגישה.")
    else:
        with st.spinner('ה-AI כותב את הפרוטוקול... זה לוקח כמה שניות'):
            try:
                # שלב 1: יצירת הטקסט
                protocol_text = generate_protocol(api_key, meeting_notes)
                
                st.success("הפרוטוקול מוכן!")
                
                # הצגה על המסך
                st.markdown("### תצוגה מקדימה:")
                st.text_area("", protocol_text, height=400)
                
                # שלב 2: יצירת הקובץ
                docx_file = create_docx(protocol_text)
                
                # כפתור הורדה
                st.download_button(
                    label="📥 הורד כקובץ Word (ערוך ושמור כ-PDF)",
                    data=docx_file.getvalue(),
                    file_name="protocol.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"אופס, משהו השתבש: {e}")
